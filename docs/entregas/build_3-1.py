"""Genera el entregable 3.1 — Desarrollo del Backend en formato .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT = Path(__file__).parent / "3.1-desarrollo-backend.docx"

doc = Document()

# Márgenes (estilo institucional)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def para(text, *, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align is not None:
        p.alignment = align
    p_format = p.paragraph_format
    p_format.space_after = Pt(8)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    return p


# ---------------------------------------------------------------- Portada
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, size, bold in [
    ("Universidad de Guadalajara", 14, True),
    ("Sistema de Universidad Virtual", 12, False),
    ("Proyecto VIII", 12, False),
    ("Unidad 3 — Actividad 3.1: Desarrollo del Backend", 13, True),
    ("Elaborado por:", 11, False),
    ("Axel Mauricio Barraza Cárdenas", 12, True),
    ("Tepic, Nayarit — 24 de mayo de 2026", 11, False),
]:
    run = title.add_run(line + "\n")
    run.font.size = Pt(size)
    run.bold = bold

doc.add_paragraph()

# ---------------------------------------------------------------- Introducción
heading("Introducción", 1)
para(
    "Este documento describe el proceso de desarrollo del backend del sistema "
    "propuesto a lo largo del semestre: una aplicación web para la tienda de ropa "
    "masculina Revolución, ubicada en el centro de Tepic, Nayarit. El sistema permite "
    "exhibir el catálogo de productos en línea, ofrecer una experiencia editorial "
    "alineada con la identidad de la marca y generar pedidos por WhatsApp que la tienda "
    "confirma y entrega de forma presencial. La propuesta, ya consolidada en las "
    "unidades anteriores, no contempla pagos en línea, cuentas de cliente ni "
    "procesamiento de checkout."
)
para(
    "La actividad integradora de la unidad 2 estableció el diseño técnico: una sola "
    "aplicación Next.js full-stack apoyada por una base de datos relacional, "
    "manteniendo el alcance acotado a un CRUD de catálogo con un panel de "
    "administración para el dueño de la tienda. Durante la implementación se "
    "respetó esa visión, pero se ajustaron tres decisiones puntuales que se "
    "documentan más adelante en este escrito: el motor de base de datos cambió de "
    "MongoDB a PostgreSQL para aprovechar Prisma y mantener un esquema fuerte; "
    "Prisma se actualizó a la versión 7 (recién publicada al iniciar el lote de "
    "trabajo); y el modelo de producto se separó en dos entidades para preparar el "
    "manejo de inventario por talla. Ninguno de estos cambios alteró el alcance "
    "funcional comprometido en la unidad 2."
)
para(
    "El trabajo se realizó dentro de la rama de Git feat/backend-prep y se ordenó en "
    "dos lotes de actividades, llamados internamente Lote A y Lote B. El Lote A "
    "comprende la infraestructura y la lectura pública del catálogo, y se encuentra "
    "completo y verificado al cierre de esta entrega. El Lote B, que cubre la "
    "autenticación y las rutas de escritura del administrador, está diseñado en "
    "detalle y se completará en paralelo con la actividad 3.2 (Desarrollo del "
    "Frontend), ya que comparte interfaz con la pantalla del panel administrativo."
)

# ---------------------------------------------------------------- Alcance
heading("Alcance del backend", 1)
para(
    "El sistema cubre tres flujos principales que ya fueron descritos en la "
    "integradora 2 y que el backend debe sostener: el cliente final navega el "
    "catálogo y arma su pedido por WhatsApp; la tienda confirma disponibilidad y "
    "coordina la recolección presencial; y el administrador gestiona los productos "
    "desde un panel privado. De estos tres flujos, el backend solo participa "
    "directamente en el primero (sirviendo el catálogo a los clientes) y en el "
    "tercero (recibiendo las altas, bajas y cambios del administrador). El segundo "
    "flujo es manual y ocurre por fuera del sistema, por lo que no requiere "
    "persistencia ni endpoints."
)
para(
    "Por esta razón, el alcance del backend se limitó a lo siguiente: un CRUD de "
    "productos con sus variantes por talla; un mecanismo de autenticación para una "
    "única cuenta administradora (la del dueño de la tienda), basado en contraseña "
    "y cookie firmada; y una superficie mínima de API REST que el frontend pueda "
    "consumir tanto en lectura pública como en escritura autenticada. Quedaron "
    "explícitamente fuera de alcance, según lo acordado en la unidad 2, el carrito "
    "persistente del lado del servidor, los pagos en línea, la bitácora de "
    "órdenes, las reservaciones de stock, la búsqueda full-text y la gestión "
    "multi-administrador. Cada una de estas piezas se documentó como diferida en "
    "una sección específica de la guía interna del proyecto, para que el siguiente "
    "ciclo de desarrollo las pueda retomar sin reabrir la discusión de alcance."
)

# ---------------------------------------------------------------- Stack
heading("Stack tecnológico definitivo", 1)
para(
    "El stack final mantiene la decisión de unificar backend y frontend en una sola "
    "aplicación Next.js, pero cambia el motor de base de datos y el ORM respecto a "
    "lo planteado inicialmente. La justificación de cada elemento se desarrolla a "
    "continuación."
)
para("Framework de aplicación", bold=True)
para(
    "Next.js 16.2.6 con App Router. Esta versión introdujo cambios importantes "
    "respecto al material disponible públicamente —por ejemplo, los parámetros de "
    "página (params y searchParams) ahora son Promesas y deben resolverse con "
    "await—, lo que obligó a leer la documentación local instalada en "
    "node_modules/next/dist/docs/ antes de escribir las primeras rutas. Los Route "
    "Handlers del App Router cubren todas las necesidades del backend de esta "
    "aplicación sin necesidad de levantar un servidor Express separado, lo que se "
    "traduce en un único repositorio, un único despliegue y un único punto de "
    "configuración."
)
para("Base de datos y ORM", bold=True)
para(
    "PostgreSQL 16 como motor de base de datos, Prisma 7.8 como ORM. Este es el "
    "ajuste técnico más importante respecto a la integradora 2, que proponía "
    "MongoDB con Mongoose. La razón del cambio es que el catálogo de Revolución "
    "tiene un esquema fuerte y bien definido (cada producto pertenece a una "
    "categoría de un conjunto cerrado, cada variante tiene una talla de un "
    "conjunto cerrado, los precios son enteros en pesos), y un motor relacional "
    "con migraciones versionadas ofrece más garantías de integridad que una base "
    "de documentos. Además, Prisma genera tipos de TypeScript a partir del "
    "esquema, lo que extiende la seguridad de tipos desde la base de datos hasta "
    "el componente que renderiza el producto, sin escribir interfaces a mano."
)
para("Validación de entrada", bold=True)
para(
    "Zod, aplicado en cada endpoint de escritura antes de tocar Prisma. La regla "
    "es estricta: ningún cuerpo de petición se confía sin validar, y los errores "
    "de validación regresan como respuestas HTTP 400 con el detalle estructurado "
    "que Zod produce. Esto evita que los errores de tipos lleguen a la base de "
    "datos como excepciones 500 y mantiene una interfaz uniforme para el frontend."
)
para("Autenticación del panel", bold=True)
para(
    "Contraseña única protegida con bcryptjs (factor 12) y sesión emitida como "
    "JWT firmado con la librería jose, almacenada en una cookie HTTP-only, "
    "Secure en producción y con SameSite=Lax. Se eligió jose en lugar de "
    "jsonwebtoken porque es compatible con el runtime edge en el que corre el "
    "middleware de Next, lo cual es necesario para proteger las rutas /admin/* "
    "y /api/admin/* sin un doble salto de servidor."
)
para("Hosting", bold=True)
para(
    "La aplicación se desplegará en Vercel y la base de datos productiva en "
    "Supabase (Postgres administrado). Para desarrollo local se utiliza un "
    "contenedor Docker con la imagen postgres:16-alpine, definido en "
    "docker-compose.yml. El esquema y las migraciones son idénticos en ambos "
    "entornos, lo que permite probar en local con plena fidelidad antes de "
    "promover a producción."
)

# ---------------------------------------------------------------- Arquitectura
heading("Arquitectura general", 1)
para(
    "El backend reside dentro del mismo proyecto Next.js y se organiza en tres "
    "capas bien delimitadas. La primera es la capa de persistencia, formada por "
    "el esquema declarado en prisma/schema.prisma y por las migraciones SQL "
    "versionadas en prisma/migrations/. La segunda es la capa de acceso a "
    "datos, encarnada por un único cliente Prisma instanciado como singleton en "
    "lib/prisma.ts; este singleton resuelve el problema clásico de Next.js, en "
    "el que el hot reload del modo desarrollo abre una conexión nueva en cada "
    "recarga y eventualmente agota el pool de conexiones de Postgres. La tercera "
    "es la capa de entrada, compuesta por los Route Handlers del App Router que "
    "viven bajo app/api/."
)
para(
    "Cada Route Handler sigue la misma estructura: recibe la petición, valida la "
    "entrada con un esquema Zod declarado en lib/schemas.ts, consulta o modifica "
    "datos a través del singleton de Prisma, y responde con un objeto JSON. Las "
    "rutas de escritura, además, invocan las funciones revalidateTag y "
    "revalidatePath de Next para que el catálogo público regenere su contenido "
    "estático en cuanto haya un cambio. Esta uniformidad facilita la lectura del "
    "código, simplifica las pruebas y reduce la cantidad de decisiones que el "
    "desarrollador debe tomar en cada nueva ruta."
)
para(
    "Las rutas se dividen en dos grupos según su nivel de privilegio. El grupo "
    "público, bajo /api/products/, expone únicamente lecturas (GET) y no requiere "
    "autenticación. El grupo administrativo, bajo /api/admin/, expone tanto "
    "escrituras como el endpoint de inicio de sesión y está protegido por un "
    "middleware global ubicado en la raíz del proyecto. Este middleware "
    "intercepta cualquier petición a /admin/* o /api/admin/* (excepto la pantalla "
    "de login), verifica la cookie de sesión y redirige al formulario de login "
    "si no es válida."
)

# ---------------------------------------------------------------- Modelo
heading("Modelo de datos", 1)
para(
    "La base de datos contiene dos tablas principales: Product y ProductVariant. "
    "La tabla Product representa cada pieza del catálogo con su slug único, "
    "nombre, descripción, precio (almacenado como entero, en pesos enteros, para "
    "evitar errores de redondeo con números de punto flotante), categoría "
    "(restringida a un enum de cinco valores: camisas, playeras, pantalones, "
    "chaquetas y accesorios), un indicador booleano de pieza destacada, y un "
    "conjunto de campos editoriales opcionales (status, edition, material, "
    "finish, lede) que reflejan el tono curatorial de la tienda. Los campos "
    "más libres (specs, shipping, story y quote) se modelaron como columnas "
    "JSONB, ya que son de solo lectura desde el frontend y no se consultan por "
    "sus campos internos: guardarlos como JSON simplifica el esquema sin perder "
    "capacidad."
)
para(
    "La tabla ProductVariant representa cada combinación producto-talla y "
    "almacena el stock disponible y una bandera isDefault que indica cuál es la "
    "talla preseleccionada en la página del producto. Esta separación es uno de "
    "los cambios más importantes respecto a la planeación inicial. La propuesta "
    "original modelaba las tallas como un arreglo dentro del producto (sizes: "
    "Size[]) y guardaba un único valor de stock por producto. Al revisar el "
    "alcance real con el dueño se notó que era previsible querer manejar stock "
    "independiente por talla en cuanto el negocio creciera, así que se hizo el "
    "rediseño ahora, antes de tener datos en producción, para evitar una "
    "migración costosa más adelante. La relación queda restringida por un "
    "índice único compuesto (productId, size) que impide duplicar variantes."
)
para(
    "El esquema completo vive en prisma/schema.prisma y se materializó en la "
    "base de datos local mediante la migración 20260525035611_init, generada "
    "con el comando npx prisma migrate dev --name init. La carpeta "
    "prisma/migrations/ está versionada en Git y es la fuente de verdad de la "
    "historia del esquema; cualquier cambio futuro pasará por una nueva "
    "migración, nunca por una edición directa del archivo SQL ya emitido."
)

# ---------------------------------------------------------------- Endpoints
heading("Superficie de la API", 1)
para(
    "La API expone dos endpoints públicos de lectura, ya implementados y "
    "verificados, y cinco endpoints administrativos planeados para el siguiente "
    "lote de trabajo. La lista completa, junto con su estado actual, es la "
    "siguiente."
)
bullet(
    "GET /api/products — listado paginado del catálogo. Acepta los parámetros "
    "de consulta page, pageSize (con un máximo de 100) y category. Devuelve un "
    "objeto con items, page, pageSize, total y totalPages, e incluye las "
    "variantes de cada producto para que el frontend pueda armar los selectores "
    "de talla sin un segundo roundtrip. Implementado y verificado."
)
bullet(
    "GET /api/products/[slug] — detalle de un producto identificado por su "
    "slug. Devuelve el producto completo con sus variantes; responde 404 "
    "(no 500) cuando el slug no existe. Implementado y verificado."
)
bullet(
    "POST /api/admin/login — recibe un cuerpo { password }, lo verifica contra "
    "ADMIN_PASSWORD_HASH y, si coincide, emite una cookie de sesión llamada "
    "revolucion-admin con duración de siete días. Planeado para el Lote B."
)
bullet(
    "POST /api/admin/logout — limpia la cookie de sesión. Planeado para el "
    "Lote B."
)
bullet(
    "POST /api/admin/products — crea un producto nuevo. Requiere sesión "
    "válida; valida el cuerpo con el esquema ProductCreate de Zod. Planeado "
    "para el Lote B."
)
bullet(
    "PATCH /api/admin/products/[slug] — actualiza parcialmente un producto. "
    "Requiere sesión válida; valida con ProductUpdate (versión .partial() de "
    "ProductCreate). Planeado para el Lote B."
)
bullet(
    "DELETE /api/admin/products/[slug] — elimina un producto y, en cascada, "
    "sus variantes. Requiere sesión válida. Planeado para el Lote B."
)
para(
    "Cada endpoint de escritura, una vez implementado, invocará revalidateTag "
    "(\"products\") y revalidatePath(\"/catalogo\") para invalidar la caché de "
    "Next y forzar la regeneración estática de la página del catálogo. Como el "
    "alcance A no expone endpoints públicos de escritura, no se requiere "
    "rate-limiting; la única superficie de escritura está detrás de "
    "autenticación."
)

# ---------------------------------------------------------------- Proceso
heading("Proceso de desarrollo", 1)
para(
    "El trabajo se ejecutó dentro de un sprint dedicado al backend y se "
    "subdividió en dos lotes secuenciales. El Lote A, completado al cierre de "
    "esta entrega, incluyó nueve pasos: levantar Postgres local en Docker, "
    "crear el proyecto en Supabase, instalar Prisma y sus dependencias, "
    "configurar las variables de entorno necesarias en .env.local, definir el "
    "esquema en prisma/schema.prisma, generar y aplicar la primera migración, "
    "construir el cliente Prisma como singleton, escribir el script de seed que "
    "puebla la base con los quince productos del catálogo legado, redactar la "
    "capa de validación con Zod y, finalmente, implementar las dos rutas "
    "públicas de lectura."
)
para(
    "El seed merece una mención particular. El proyecto venía operando con un "
    "arreglo hand-edited de productos exportado desde lib/products.ts, que el "
    "frontend leía directamente. En lugar de descartar ese arreglo se "
    "reutilizó como fuente de datos del seed: el script prisma/seed.ts itera "
    "sobre los quince productos, llama a prisma.product.upsert para no "
    "duplicar registros si se vuelve a correr, y crea las variantes "
    "correspondientes a cada talla. Esto permitió tener datos realistas en la "
    "base desde el primer día y se conservó el arreglo legado como fallback "
    "temporal mientras el frontend se migra a leer desde la base de datos en "
    "la siguiente actividad."
)
para(
    "El Lote B, planeado pero pendiente de implementación, cubrirá la "
    "autenticación y las rutas de escritura. Concretamente, se debe escribir "
    "lib/auth.ts con las funciones verifyPassword, signSession y "
    "verifySession; crear middleware.ts en la raíz del proyecto con el matcher "
    "que protege /admin/((?!login).*) y /api/admin/:path*; generar el hash de "
    "contraseña y el secreto de sesión, y guardarlos en ADMIN_PASSWORD_HASH y "
    "ADMIN_SESSION_SECRET de .env.local; e implementar las cinco rutas "
    "administrativas. Este lote se completará en paralelo con la actividad "
    "3.2, ya que comparte interfaz con la pantalla de login y con el listado "
    "editable de productos del panel administrativo."
)

# ---------------------------------------------------------------- Dificultades
heading("Dificultades encontradas y cómo se resolvieron", 1)
para(
    "El proceso no fue lineal. Tres problemas significativos obligaron a "
    "desviarse del plan original y se documentan a continuación con el detalle "
    "suficiente para que el siguiente desarrollador los reconozca y los "
    "resuelva más rápido."
)
para("Conflicto del puerto 5432 con el servicio de Windows", bold=True)
para(
    "Al ejecutar por primera vez docker compose up -d el contenedor se "
    "levantaba aparentemente sin errores, pero las consultas a "
    "localhost:5432 caían en una instalación previa de Postgres 17 que tenía "
    "Windows registrada como servicio (postgresql-x64-17). Docker Desktop no "
    "muestra advertencias claras cuando un bind de puerto falla "
    "silenciosamente en Windows, lo que dificultó el diagnóstico: el comando "
    "docker compose ps reportaba el contenedor como running y con el puerto "
    "mapeado, pero las herramientas del host se conectaban al servicio "
    "equivocado. La solución fue remapear el puerto del contenedor de "
    "5432:5432 a 5433:5432 en docker-compose.yml y actualizar DATABASE_URL "
    "a localhost:5433. El cambio quedó documentado en la guía interna del "
    "proyecto (docs/backend-setup.md) y, si en algún momento se detiene el "
    "servicio de Windows para recuperar el puerto canónico, basta con "
    "revertir ambos archivos juntos."
)
para("Cambios breaking de Prisma 7 respecto a la guía v6", bold=True)
para(
    "La guía interna del proyecto se había escrito esperando Prisma 6, pero "
    "al instalar las dependencias se descargó la versión 7.8 (recién "
    "publicada al iniciar el sprint). Esta versión introdujo varios cambios "
    "incompatibles que rompieron cinco suposiciones del documento. Primero, "
    "el generador del cliente ahora se llama prisma-client en lugar de "
    "prisma-client-js y se debe declarar una ruta de salida explícita; el "
    "proyecto la apunta a ../lib/generated/prisma, y todas las importaciones "
    "se hacen desde @/lib/generated/prisma/client. Segundo, los campos url "
    "y directUrl ya no se permiten dentro de schema.prisma: la URL debe "
    "vivir en el nuevo archivo prisma.config.ts. Tercero, prisma.config.ts "
    "no lee .env.local automáticamente, por lo que fue necesario importar "
    "dotenv y llamar a loadEnv({ path: \".env.local\" }) al inicio del "
    "archivo. Cuarto, el comando prisma migrate dev ya no ejecuta prisma "
    "generate de manera implícita; ahora hay que invocarlo como un paso "
    "explícito. Quinto, el constructor new PrismaClient() ya no acepta una "
    "llamada sin argumentos: exige un driver adapter, lo que obligó a "
    "instalar @prisma/adapter-pg junto con pg y a construir el cliente como "
    "new PrismaClient({ adapter: new PrismaPg(DATABASE_URL) })."
)
para(
    "Estos cambios consumieron una tarde completa de depuración. Para "
    "evitar que vuelvan a sorprender, se agregó a la guía interna una "
    "sección titulada Prisma v7 deviations que enumera cada una y muestra "
    "el código tal como quedó después del ajuste. La sección sirve a la vez "
    "como bitácora del incidente y como referencia para cualquier desarrollo "
    "futuro contra Prisma 7."
)
para("Columnas JSON y tipos estrictos de TypeScript", bold=True)
para(
    "El script de seed comenzó a fallar al pasar los campos editoriales "
    "(specs, shipping, story y quote) al método prisma.product.create. El "
    "compilador de TypeScript marcaba un error TS2322, alegando que las "
    "interfaces tipadas de esos campos no eran asignables a "
    "Prisma.InputJsonValue. La causa de fondo es que InputJsonValue requiere "
    "una index signature ({ [key: string]: ... }), y las interfaces tipadas "
    "del proyecto deliberadamente no la tienen para conservar autocompletado "
    "y verificación estricta en el frontend. La solución fue introducir un "
    "helper genérico asJson<T>() que realiza una conversión controlada a "
    "través de unknown hacia InputJsonValue. De este modo los tipos "
    "estrictos siguen rigiendo en el resto del código y el ensanchamiento "
    "ocurre solo en la frontera con Prisma. El mismo patrón se reaplicará "
    "en las rutas de escritura del Lote B."
)

# ---------------------------------------------------------------- Cambios
heading("Cambios respecto a la planeación de la unidad 2", 1)
para(
    "La integradora 2 fijó las decisiones de arquitectura y alcance, pero "
    "como cualquier planeación previa a la construcción hubo elementos que "
    "se ajustaron al chocar con la realidad técnica. Los cambios fueron los "
    "siguientes."
)
para(
    "El motor de base de datos pasó de MongoDB con Mongoose a PostgreSQL con "
    "Prisma 7.8. La razón principal es que el catálogo tiene un esquema "
    "fuerte y bien delimitado (cinco categorías, cinco tallas, precios "
    "enteros) y un motor relacional con migraciones versionadas ofrece más "
    "garantías de integridad. Además, Prisma genera los tipos de "
    "TypeScript a partir del esquema, lo que cierra el ciclo de seguridad "
    "de tipos desde la base de datos hasta el componente React, sin "
    "necesidad de mantener interfaces a mano."
)
para(
    "El modelo de producto se separó en dos entidades (Product y "
    "ProductVariant) en lugar del modelo plano original con un arreglo de "
    "tallas y un único campo de stock. Como ya se explicó, esta separación "
    "permite manejar inventario por talla sin migraciones futuras y refleja "
    "mejor la realidad del negocio."
)
para(
    "Las dependencias técnicas se ajustaron a las versiones disponibles al "
    "instalar: Next.js quedó en 16.2.6 y Prisma en 7.8. Estos cambios "
    "no estaban planeados pero tampoco contradicen el espíritu de la "
    "propuesta original; simplemente reflejan que el ecosistema evoluciona "
    "más rápido que un documento de diseño."
)
para(
    "El puerto del Postgres local pasó de 5432 a 5433 por el conflicto con "
    "el servicio de Windows, según ya se explicó."
)
para(
    "Finalmente, la configuración del seed se movió de package.json (donde "
    "Prisma 6 la esperaba) a prisma.config.ts (donde Prisma 7 la exige). "
    "Es un cambio menor pero crítico: sin él, npx prisma db seed responde "
    "con el mensaje No seed command configured y no ejecuta nada."
)
para(
    "Ninguno de estos ajustes alteró el alcance funcional comprometido en "
    "la unidad 2. El sistema sigue siendo un CRUD de catálogo con "
    "autenticación de un solo administrador y un flujo de venta que termina "
    "en WhatsApp; lo que cambió fueron los detalles de implementación, no "
    "las funcionalidades visibles para el usuario."
)

# ---------------------------------------------------------------- Verificaciones
heading("Verificaciones realizadas", 1)
para(
    "Al cierre del Lote A se ejecutó la lista de verificación definida en "
    "la guía interna del proyecto. Los resultados se resumen a continuación, "
    "marcando con un guion los pasos pendientes que corresponden al Lote B."
)
bullet(
    "docker compose up -d levanta Postgres y docker compose ps lo reporta "
    "como running en el puerto 5433. Verificado."
)
bullet(
    "npx prisma migrate status reporta la migración 20260525035611_init "
    "aplicada al día. Verificado."
)
bullet(
    "npx prisma db seed siembra los quince productos del catálogo legado "
    "sin errores. Verificado."
)
bullet(
    "npx prisma studio abre el navegador y muestra todos los productos con "
    "sus variantes por talla. Verificado."
)
bullet(
    "curl http://localhost:3000/api/products regresa un JSON paginado con "
    "los productos sembrados. Verificado."
)
bullet(
    "curl http://localhost:3000/api/products/<slug> regresa un producto "
    "individual con sus variantes; un slug inexistente regresa 404. "
    "Verificado."
)
bullet(
    "curl http://localhost:3000/api/products?page=-1 regresa 400 con el "
    "detalle estructurado de Zod, no 500. Verificado."
)
bullet(
    "POST /api/admin/products sin cookie de sesión debe regresar 401 o "
    "redirigir a /admin/login. Pendiente (Lote B)."
)
bullet(
    "POST /api/admin/login con la contraseña correcta debe emitir la "
    "cookie de sesión, y la siguiente petición a POST /api/admin/products "
    "debe tener éxito. Pendiente (Lote B)."
)
bullet(
    "yarn build debe completar sin errores de TypeScript ni desviaciones "
    "del tipo de Prisma. Por re-ejecutar al cerrar el Lote B."
)

# ---------------------------------------------------------------- Capturas
heading("Capturas de pantalla", 1)
para(
    "Las siguientes capturas evidencian cada paso del proceso. Las imágenes "
    "se entregan adjuntas al buzón de la actividad con los nombres de "
    "archivo indicados al pie de cada figura."
)

figs = [
    (
        "Figura 1. Contenedor de Postgres corriendo",
        "01-docker-ps.png",
        "Salida del comando docker compose ps. El contenedor revolucion-db "
        "(imagen postgres:16-alpine) aparece en estado running y mapea el "
        "puerto interno 5432 al puerto 5433 del host, ajuste necesario por "
        "el conflicto con el servicio de Windows postgresql-x64-17.",
    ),
    (
        "Figura 2. Esquema de Prisma",
        "02-schema-prisma.png",
        "Vista de prisma/schema.prisma en el editor. Se observan los dos "
        "modelos (Product y ProductVariant), los enums Category y Size, y "
        "el datasource db sin el campo url —Prisma 7 prohíbe declararlo "
        "ahí; ahora vive en prisma.config.ts.",
    ),
    (
        "Figura 3. Migración aplicada",
        "03-prisma-migrate.png",
        "Ejecución de npx prisma migrate dev --name init. Prisma crea la "
        "carpeta prisma/migrations/20260525035611_init/, escribe el SQL de "
        "la migración, lo aplica contra la base local y deja el esquema "
        "en sincronía.",
    ),
    (
        "Figura 4. Cliente Prisma generado",
        "04-prisma-generate.png",
        "Resultado de npx prisma generate. El cliente se emite a "
        "lib/generated/prisma/ (configurado en el campo output del "
        "generador) y, desde el código de la aplicación, se importa como "
        "@/lib/generated/prisma/client.",
    ),
    (
        "Figura 5. Seed exitoso",
        "05-seed.png",
        "Salida de npx prisma db seed. El script prisma/seed.ts corre con "
        "tsx, hace upsert de los quince productos del catálogo legado y "
        "crea sus variantes por talla. El mensaje final confirma "
        "\"Seeded 15 products.\"",
    ),
    (
        "Figura 6. Datos verificados en Prisma Studio",
        "06-prisma-studio.png",
        "Vista de la tabla Product en npx prisma studio (puerto 5555). "
        "Los quince registros sembrados aparecen con todos sus campos "
        "editoriales y enlazados a sus ProductVariant correspondientes.",
    ),
    (
        "Figura 7. Endpoint público GET /api/products",
        "07-api-products.png",
        "Respuesta del endpoint público de listado. El JSON tiene la "
        "forma { items, page, pageSize, total, totalPages } y cada item "
        "trae sus variantes incluidas, de modo que el frontend pueda armar "
        "los selectores de talla en un solo roundtrip.",
    ),
    (
        "Figura 8. Endpoint GET /api/products/[slug]",
        "08-api-product-slug.png",
        "Respuesta del endpoint de detalle. Devuelve un producto completo "
        "con sus variantes y todos los campos editoriales (specs, story, "
        "quote). Cuando el slug no existe, el endpoint responde 404 en "
        "lugar de 500.",
    ),
    (
        "Figura 9. Validación con Zod (entrada inválida)",
        "09-zod-400.png",
        "Caso de prueba con un parámetro de consulta inválido (page=-1). "
        "El esquema ProductListQuery rechaza la entrada antes de tocar "
        "Prisma y responde con HTTP 400 más el detalle estructurado de "
        "Zod. Ningún error de validación se filtra como excepción 500.",
    ),
    (
        "Figura 10. Estructura final del repositorio",
        "10-tree.png",
        "Vista general del repositorio después del Lote A. Los archivos "
        "nuevos del backend (prisma/, prisma.config.ts, docker-compose.yml, "
        "lib/prisma.ts, lib/schemas.ts y app/api/products/*) conviven con "
        "el frontend ya existente sin modificarlo.",
    ),
]

for title_text, filename, desc in figs:
    p = doc.add_paragraph()
    run = p.add_run(f"[ {title_text} — adjunto: {filename} ]")
    run.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(f"{title_text}. {desc}")

# ---------------------------------------------------------------- Conclusión
heading("Conclusión", 1)
para(
    "El backend cumple con el alcance acordado en la unidad 2 en su parte "
    "de lectura pública: base de datos modelada, migrada y poblada con datos "
    "realistas; cliente Prisma robusto contra el hot reload de Next; capa de "
    "validación con Zod aplicada antes de cualquier consulta; rutas GET "
    "paginadas y verificadas con curl; y una guía interna que documenta "
    "cada decisión y cada desviación. La arquitectura unificada Next.js "
    "full-stack, propuesta originalmente en la integradora 2, se sostuvo "
    "sin sobresaltos: un solo repositorio, una sola base de código, un solo "
    "despliegue."
)
para(
    "El Lote B (autenticación y escritura administrativa) está diseñado en "
    "detalle y a un par de horas de implementación; se cerrará en paralelo "
    "con la actividad 3.2, donde comparte interfaz con la pantalla del "
    "panel administrativo. Las dificultades técnicas más costosas —el "
    "conflicto del puerto 5432 con Windows y los cambios breaking de "
    "Prisma 7— se resolvieron sin reducir el alcance: solo obligaron a "
    "ajustar detalles de implementación. La planeación original en cuanto "
    "a modelos, endpoints y flujo de venta vía WhatsApp permaneció "
    "intacta."
)
para(
    "El código fuente y el avance funcional pueden consultarse en el "
    "repositorio del proyecto. El README documenta los comandos de "
    "instalación, las variables de entorno requeridas y el estado actual "
    "de cada lote de trabajo."
)
para(
    "Repositorio: https://github.com/AxelMBC/tienda-revolucion (rama "
    "feat/backend-prep)."
)

# ---------------------------------------------------------------- Referencias
heading("Referencias", 1)
para(
    "Altman, H. (2018). Scrum: La Primera Metodología Ágil para Gestionar "
    "el Desarrollo de Productos Paso a Paso. CreateSpace."
)
para(
    "Edge, J. (2020). Agile: Una guía esencial sobre la gestión de "
    "proyectos Agile, el proceso Kanban y el Lean Thinking. Bravex "
    "Publications."
)
para(
    "Cáceres Yujra, R. (2021). Uso de la metodología Scrum en el proceso "
    "de pruebas de software del desarrollo de una aplicación móvil del "
    "sector bancario [Tesis de pregrado, Universidad Nacional Mayor de "
    "San Marcos]. https://hdl.handle.net/20.500.12672/17840"
)
para(
    "Gómez Trapero, J. (2016). Desarrollo front-end y back-end de una "
    "aplicación Android distribuida para el sector de la restauración "
    "[Tesis de pregrado, Universidad Politécnica de Madrid]. "
    "https://oa.upm.es/43493/"
)
para(
    "Bedoya Alzate, E. (2021). Implementación de pruebas unitarias "
    "[Tesis de pregrado, Universidad de Antioquía]. "
    "https://hdl.handle.net/10495/20225"
)

doc.save(OUT)
print(f"Documento generado: {OUT}")
