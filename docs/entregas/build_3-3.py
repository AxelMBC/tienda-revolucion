"""Genera el entregable 3.3 — Implementación de los sistemas Backend y Frontend en formato .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT = Path(__file__).parent / "3.3-implementacion-sistemas.docx"

doc = Document()

# Márgenes (estilo institucional)
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def para(text, *, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    return p


# ---------------------------------------------------------------- Portada
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line, size, bold in [
    ("Universidad de Guadalajara", 14, True),
    ("Sistema de Universidad Virtual", 12, False),
    ("Proyecto VIII — IH740_2026A", 12, False),
    ("Unidad 3 — Actividad 3.3: Implementación de los sistemas Backend y Frontend", 13, True),
    ("Elaborado por:", 11, False),
    ("Axel Mauricio Barraza Cárdenas", 12, True),
    ("Tepic, Nayarit — 24 de mayo de 2026", 11, False),
]:
    run = title.add_run(line + "\n")
    run.font.size = Pt(size)
    run.bold = bold

doc.add_paragraph()

# ---------------------------------------------------------------- Introducción
heading("Introducción", 1)
para(
    "Este reporte integrador describe la implementación de los sistemas Backend y "
    "Frontend del proyecto desarrollado durante el semestre: el sitio web de la Tienda "
    "de Ropa Revolución, una tienda de ropa para hombre en el centro de Tepic, "
    "Nayarit. En las actividades previas de la unidad se construyó el backend "
    "(actividad 3.1) —PostgreSQL, ORM Prisma y una API REST sobre Next.js— y se "
    "desarrolló el frontend (actividad 3.2) con el sistema de diseño de la marca. La "
    "actividad 3.3 cierra el ciclo: poner ambos sistemas a funcionar juntos y dejarlos "
    "listos para operar en un entorno real."
)
para(
    "Como advierte la actividad, una parte central de la implementación es decidir "
    "cómo y dónde correrá el sistema, y reconocer que el entorno donde se construye no "
    "es el mismo donde vivirá en producción. Este proyecto enfrenta esa situación: el "
    "desarrollo ocurre en una computadora propia, contra una base de datos PostgreSQL "
    "local en Docker; pero el destino productivo es distinto —Vercel para la "
    "aplicación y Supabase para la base de datos administrada—, con diferencias que "
    "impactan al sistema y que se documentan aquí. El trabajo continuó en la rama de "
    "Git feat/backend-prep. Al cierre de esta entrega, la implementación local está "
    "completa y verificada (el frontend ya lee de la base de datos, no de un arreglo "
    "escrito a mano), y el entorno de producción está aprovisionado con su "
    "procedimiento de despliegue documentado; la promoción final es el paso pendiente."
)

# ---------------------------------------------------------------- Qué significa implementar
heading("Qué significa “implementar” en este proyecto", 1)
para(
    "La implementación abarca tres trabajos concretos: (1) conectar el frontend al "
    "backend para que las páginas consuman datos reales de la base de datos en lugar "
    "del catálogo escrito a mano; (2) levantar y operar el sistema completo en el "
    "entorno de desarrollo local; y (3) preparar y documentar el despliegue al entorno "
    "de producción. El sistema es una aplicación Next.js full-stack —backend y "
    "frontend viven en el mismo proyecto y se despliegan como una sola unidad—, lo que "
    "simplifica la operación pero no elimina la distinción entre capas: el backend es "
    "la base de datos, el cliente Prisma y las rutas de API; el frontend son las "
    "páginas y componentes React. La frontera entre ambos, y la pieza clave de esta "
    "actividad, es la capa de acceso a datos lib/catalog.ts."
)

# ---------------------------------------------------------------- Entorno local
heading("Entorno de implementación local", 1)
para(
    "El primer entorno donde se implementó el sistema es la computadora de desarrollo "
    "(Windows 11). Sus componentes son los siguientes."
)
para(
    "PostgreSQL 16 corre en un contenedor definido en docker-compose.yml con la imagen "
    "postgres:16-alpine y un volumen nombrado para persistir los datos. El contenedor "
    "expone la base en el puerto 5433 del host (no el 5432 habitual) por un conflicto "
    "con un servicio de Postgres ya instalado en Windows, según se detalla en la "
    "sección de dificultades; se levanta con docker compose up -d y se verifica con "
    "docker compose ps."
)
para(
    "La aplicación se ejecuta con yarn dev en localhost:3000 y lee su configuración "
    "desde .env.local, donde DATABASE_URL y DIRECT_URL apuntan al contenedor en "
    "localhost:5433. En este entorno ambas URL son idénticas, porque no hay un pooler "
    "de conexiones intermediando entre la aplicación y Postgres. La base se pobló con "
    "el script prisma/seed.ts, que reutiliza el catálogo legado de lib/products.ts "
    "(quince productos) y los inserta con upsert junto con sus variantes por talla, "
    "dando datos realistas e idénticos a los que el sitio mostraba antes para comparar "
    "el comportamiento de extremo a extremo."
)

# ---------------------------------------------------------------- Implementación del frontend
heading("Implementación: conexión del Frontend con el Backend", 1)
para(
    "El núcleo de esta actividad fue conectar el frontend ya existente al backend. "
    "Hasta este punto, las páginas del sitio importaban sus datos directamente del "
    "arreglo escrito a mano en lib/products.ts. La implementación consistió en "
    "interponer una capa de acceso a datos que conserva exactamente la misma interfaz "
    "de consulta, pero que por dentro lee de la base de datos a través de Prisma."
)
para("La capa adaptadora lib/catalog.ts", bold=True)
para(
    "Se creó el módulo lib/catalog.ts como nuevo origen de datos del sitio. Expone las "
    "mismas funciones de consulta que el frontend ya usaba (getAllProducts, "
    "getFeatured, getPicks, getBySlug, getByCategory, filterProducts, "
    "getRelatedPieces y getCategoryCounts), pero cada una ejecuta ahora una consulta "
    "Prisma contra PostgreSQL en lugar de filtrar un arreglo en memoria. La pieza que "
    "hace transparente el cambio es la función interna mapDbProduct, que traduce cada "
    "fila de la base de datos (modelos Product y ProductVariant) al tipo Product que "
    "los componentes ya esperaban, definido en lib/types.ts: las variantes por talla "
    "se colapsan en los campos sizes, defaultSize y stock, y los campos editoriales en "
    "JSON (specs, shipping, story, quote) se reconstruyen con su forma tipada. Gracias "
    "a ello, los componentes de presentación no se reescribieron —siguen recibiendo "
    "objetos Product con la misma forma—; el cambio se concentró en el origen de los "
    "datos, no en cómo se muestran."
)
para("Páginas conectadas y asincronía", bold=True)
para(
    "Las cuatro superficies que muestran catálogo se repuntaron para leer del nuevo "
    "módulo: catálogo (app/catalogo/page.tsx), ofertas (app/ofertas/page.tsx), detalle "
    "de producto (app/producto/[slug]/page.tsx) y las piezas destacadas de la portada "
    "(components/home/ThePieces.tsx). El cambio principal fue pasar de datos síncronos "
    "a esperar (await) las nuevas funciones, que devuelven Promesas porque consultan "
    "la base de datos. Esto encaja con Next.js 16, donde los componentes de servidor "
    "pueden ser asíncronos y los parámetros de página (params, searchParams) ya son "
    "Promesas que se resuelven con await. Se conservó lib/products.ts como respaldo y "
    "fuente del seed; su retiro definitivo ocurrirá cuando el panel de administración "
    "(pendiente) permita editar el catálogo directamente en la base."
)

# ---------------------------------------------------------------- Entorno de producción
heading("Entorno de producción: cómo y dónde", 1)
para(
    "El entorno donde el sistema vivirá en la vida real es distinto del entorno de "
    "desarrollo, y esa diferencia es justamente lo que la actividad pide visualizar. "
    "La aplicación se desplegará en Vercel (la plataforma natural para Next.js, con "
    "ejecución serverless) y la base de datos productiva será Supabase, que ofrece "
    "PostgreSQL administrado. No se usa una máquina virtual propia ni un servidor "
    "dedicado: se eligió un modelo administrado para reducir la carga de operación de "
    "una tienda pequeña."
)
para(
    "El esquema y las migraciones son idénticos en local y en producción, de modo que "
    "la base administrada de Supabase tendrá exactamente la misma estructura que la "
    "base local de Docker. Lo que cambia no es el modelo de datos, sino la forma de "
    "conectarse a él y de ejecutar la aplicación, que es donde aparecen las "
    "diferencias relevantes."
)
para("Diferencias entre el entorno inicial y el entorno real", bold=True)
para(
    "Llevar el sistema del entorno local al productivo no es copiar archivos; varias "
    "diferencias de entorno impactan directamente al sistema y obligan a configurarlo "
    "distinto:"
)
bullet(
    "Dos cadenas de conexión, no una. En producción, Supabase expone una conexión "
    "agrupada (pooler PgBouncer, puerto 6543) que las funciones serverless deben usar "
    "para no agotar las conexiones de Postgres, y una conexión directa (puerto 5432) "
    "que solo se usa para aplicar migraciones. En local existe una sola base en "
    "Docker, así que DATABASE_URL y DIRECT_URL coinciden; en producción son "
    "diferentes."
)
bullet(
    "Parámetros obligatorios del pooler. La URL agrupada de producción requiere los "
    "parámetros ?pgbouncer=true&connection_limit=1 para que Prisma funcione "
    "correctamente detrás de PgBouncer en modo transacción. En local no aplican."
)
bullet(
    "Las variables de entorno no viven en un archivo. En local están en .env.local; "
    "en producción se configuran en el panel de Vercel y nunca se suben al "
    "repositorio. Los secretos (hash de contraseña del administrador y secreto de "
    "sesión) se generan una sola vez y se guardan solo ahí."
)
bullet(
    "Migraciones con migrate deploy, no migrate dev. En producción jamás se ejecuta "
    "el comando interactivo de desarrollo; el esquema se aplica con prisma migrate "
    "deploy, que solo corre las migraciones ya versionadas, usando la conexión "
    "directa."
)
bullet(
    "Generación del cliente en cada despliegue. El cliente Prisma se regenera "
    "automáticamente en producción mediante el script postinstall (prisma generate) "
    "del package.json, porque la carpeta lib/generated/prisma no se sube al "
    "repositorio."
)

# ---------------------------------------------------------------- Proceso de despliegue
heading("Procedimiento de despliegue a producción", 1)
para(
    "El despliegue al entorno real se documentó como una secuencia reproducible en la "
    "guía interna del proyecto (docs/backend-setup.md). Los pasos son los siguientes:"
)
bullet(
    "Crear el proyecto en Supabase eligiendo la región más cercana a los usuarios "
    "(us-west-1 para Tepic) y resguardar la contraseña de la base. Hecho."
)
bullet(
    "Obtener del panel de Supabase las dos cadenas de conexión —pooler (6543) y "
    "directa (5432)— para usarlas como DATABASE_URL y DIRECT_URL en producción."
)
bullet(
    "Importar el repositorio en Vercel y capturar las cuatro variables de entorno del "
    "backend (DATABASE_URL, DIRECT_URL, ADMIN_PASSWORD_HASH, ADMIN_SESSION_SECRET) "
    "más las del frontend (NEXT_PUBLIC_WHATSAPP_PHONE, NEXT_PUBLIC_SITE_URL) en el "
    "panel de Vercel."
)
bullet(
    "Aplicar el esquema a Supabase con prisma migrate deploy usando la conexión "
    "directa, y ejecutar el seed una sola vez para poblar el catálogo inicial."
)
bullet(
    "Disparar el build de producción en Vercel (next build), que regenera el cliente "
    "Prisma vía postinstall y compila la aplicación; verificar que el sitio en línea "
    "muestra el catálogo leído desde Supabase."
)
para(
    "Al cierre de esta entrega, el proyecto de Supabase ya existe y el procedimiento "
    "está documentado y listo para ejecutarse; la promoción final (cargar las URL en "
    "Vercel, correr migrate deploy y el seed contra Supabase, y lanzar el primer "
    "build productivo) es el paso pendiente. Se mantiene como último porque conviene "
    "ejecutarlo una sola vez, con todas las variables ya verificadas, para no dejar "
    "estados intermedios en la base productiva."
)

# ---------------------------------------------------------------- Dificultades
heading("Dificultades encontradas durante la implementación", 1)
para(
    "La implementación expuso problemas que no aparecen durante el desarrollo aislado "
    "de cada capa y que confirman la advertencia de la actividad: el entorno importa."
)
para("Conflicto de puerto con el Postgres de Windows", bold=True)
para(
    "Al levantar el contenedor, las conexiones a localhost:5432 caían en una "
    "instalación previa de Postgres registrada como servicio de Windows "
    "(postgresql-x64-17), no en el contenedor. Docker Desktop no advierte con claridad "
    "cuando un mapeo de puerto falla en silencio en Windows: docker compose ps "
    "reportaba el contenedor corriendo y con el puerto mapeado, pero el host hablaba "
    "con el servicio equivocado. Se resolvió remapeando el puerto a 5433:5432 en "
    "docker-compose.yml y actualizando las URL de .env.local a localhost:5433."
)
para("Conexiones de base de datos en un entorno serverless", bold=True)
para(
    "La mayor diferencia entre local y producción es el manejo de conexiones. En "
    "desarrollo, el cliente Prisma se construye como singleton en lib/prisma.ts para "
    "evitar que el hot reload de Next abra una conexión nueva en cada recarga y agote "
    "el pool. En producción cada invocación serverless en Vercel puede abrir su propia "
    "conexión y saturar la base; por eso Supabase obliga a usar el pooler PgBouncer "
    "(puerto 6543) con connection_limit=1, mientras que las migraciones usan la "
    "conexión directa (puerto 5432). Anticipar esto evitó un error clásico —desplegar "
    "apuntando a la conexión directa y ver caer el sitio bajo tráfico— y motivó la "
    "separación de DATABASE_URL y DIRECT_URL."
)
para("Columnas JSON y tipos estrictos al mapear de la base al frontend", bold=True)
para(
    "Los campos editoriales (specs, shipping, story, quote) se guardan como JSON, pero "
    "el frontend los consume con interfaces tipadas estrictas, y al leerlos de Prisma "
    "llegan como un JSON genérico no asignable a esas interfaces. En mapDbProduct se "
    "resolvió con una conversión controlada a través de unknown hacia el tipo "
    "esperado, conservando la seguridad de tipos en los componentes y limitando el "
    "ensanchamiento a la frontera con la base de datos."
)

# ---------------------------------------------------------------- Verificaciones
heading("Verificaciones realizadas", 1)
para(
    "La implementación en el entorno local se validó con la siguiente lista. Los "
    "pasos del entorno de producción quedan marcados como pendientes de la promoción "
    "final."
)
bullet(
    "docker compose up -d levanta Postgres y docker compose ps lo reporta como "
    "running en el puerto 5433. Verificado."
)
bullet(
    "npx prisma migrate status reporta el esquema al día y npx prisma db seed siembra "
    "los quince productos sin errores. Verificado."
)
bullet(
    "La página de catálogo (/catalogo) renderiza los productos leídos desde la base "
    "de datos vía lib/catalog.ts, con sus filtros por categoría, talla y orden. "
    "Verificado."
)
bullet(
    "La página de detalle (/producto/[slug]) muestra un producto con sus variantes, "
    "sus campos editoriales y sus piezas relacionadas; un slug inexistente produce la "
    "página 404. Verificado."
)
bullet(
    "La portada y la página de ofertas muestran las piezas destacadas obtenidas con "
    "getFeatured/getPicks desde la base de datos. Verificado."
)
bullet(
    "Los endpoints públicos GET /api/products (paginado, con validación Zod) y GET "
    "/api/products/[slug] responden con los datos sembrados. Verificado."
)
bullet(
    "prisma migrate deploy y el seed ejecutados contra Supabase con las variables de "
    "producción. Pendiente (promoción final)."
)
bullet(
    "Sitio en línea servido por Vercel mostrando el catálogo leído desde Supabase; "
    "yarn build sin errores de TypeScript de extremo a extremo. Pendiente (promoción "
    "final)."
)

# ---------------------------------------------------------------- Capturas
heading("Capturas de pantalla", 1)
para(
    "Las siguientes capturas evidencian el proceso de implementación. Se entregan "
    "adjuntas al buzón de la actividad; guárdense en docs/entregas/img/3-3/ con los "
    "nombres indicados al pie de cada figura."
)

figs = [
    (
        "Figura 1. Base de datos local en operación",
        "01-docker-ps.png",
        "Salida de docker compose ps. El contenedor revolucion-db "
        "(postgres:16-alpine) aparece en estado running, mapeando el puerto interno "
        "5432 al 5433 del host por el conflicto con el servicio de Windows.",
    ),
    (
        "Figura 2. Frontend leyendo del backend",
        "02-catalogo-db.png",
        "Página /catalogo en localhost:3000 renderizando los productos obtenidos de "
        "PostgreSQL a través de lib/catalog.ts. Los datos provienen de la base, no del "
        "arreglo lib/products.ts.",
    ),
    (
        "Figura 3. Página de detalle con datos reales",
        "03-producto-slug.png",
        "Página /producto/[slug] mostrando un producto con sus variantes por talla y "
        "sus campos editoriales, todo leído de la base de datos vía getBySlug.",
    ),
    (
        "Figura 4. Proyecto de Supabase (entorno productivo)",
        "04-supabase.png",
        "Panel de Supabase con el proyecto creado y las dos cadenas de conexión: "
        "pooler en el puerto 6543 (runtime serverless) y directa en el 5432 "
        "(migraciones).",
    ),
    (
        "Figura 5. Variables de entorno en Vercel",
        "05-vercel-env.png",
        "Panel de variables de entorno de Vercel con las claves del backend "
        "(DATABASE_URL, DIRECT_URL, ADMIN_PASSWORD_HASH, ADMIN_SESSION_SECRET) y del "
        "frontend, capturadas fuera del repositorio.",
    ),
]

for title_text, filename, desc in figs:
    p = doc.add_paragraph()
    run = p.add_run(f"[ {title_text} — adjunto: {filename} ]")
    run.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(f"{title_text}. {desc}")

# ---------------------------------------------------------------- Estado y pendientes
heading("Estado actual y trabajo pendiente", 1)
para(
    "El sistema está implementado y verificado en el entorno local: el frontend lee "
    "del backend a través de la base de datos, no de un arreglo estático, y ambas "
    "capas operan juntas de extremo a extremo. El entorno de producción está "
    "aprovisionado (proyecto de Supabase creado) y su despliegue documentado y listo "
    "para ejecutarse."
)
para(
    "Queda pendiente la promoción final a producción —cargar las variables en Vercel, "
    "aplicar las migraciones y el seed contra Supabase y lanzar el primer build "
    "productivo— y, en paralelo, el panel de administración con autenticación, que "
    "permitirá al dueño de la tienda gestionar el catálogo directamente sobre la base "
    "de datos y retirar definitivamente el arreglo legado. Esa funcionalidad de "
    "escritura se valida en la actividad 3.4 (pruebas al sistema)."
)

# ---------------------------------------------------------------- Conclusión
heading("Conclusión", 1)
para(
    "La implementación confirmó la premisa de la actividad: el entorno donde se "
    "construye un sistema no es el mismo donde vivirá, y prever esa transición es "
    "parte del trabajo. La integración del frontend con el backend se resolvió de "
    "forma limpia gracias a una capa adaptadora (lib/catalog.ts) que mantuvo intacta "
    "la interfaz que los componentes esperaban, de modo que cambió el origen de los "
    "datos sin reescribir la presentación. En el entorno local el sistema quedó "
    "completo y verificado."
)
para(
    "Las diferencias entre el entorno inicial (PostgreSQL local en Docker, una sola "
    "conexión, variables en archivo) y el entorno real (Supabase administrado con "
    "pooler de conexiones, conexión directa para migraciones, variables en el panel "
    "de Vercel, ejecución serverless) se identificaron, documentaron y configuraron "
    "por anticipado, lo que reduce el riesgo de la promoción final a una secuencia "
    "reproducible. El desarrollo se documentó paso a paso y cada incidencia quedó "
    "registrada en la guía interna del proyecto, cumpliendo la recomendación de "
    "documentar el proceso y cualquier incidencia."
)
para(
    "El código fuente y el avance funcional pueden consultarse en el repositorio del "
    "proyecto, rama feat/backend-prep: "
    "https://github.com/AxelMBC/tienda-revolucion"
)

# ---------------------------------------------------------------- Referencias
heading("Referencias", 1)
para(
    "Altman, H. (2018). Scrum: La Primera Metodología Ágil para Gestionar el "
    "Desarrollo de Productos Paso a Paso. CreateSpace."
)
para(
    "Edge, J. (2020). Agile: Una guía esencial sobre la gestión de proyectos Agile, "
    "el proceso Kanban y el Lean Thinking. Bravex Publications."
)
para(
    "Cáceres Yujra, R. (2021). Uso de la metodología Scrum en el proceso de pruebas "
    "de software del desarrollo de una aplicación móvil del sector bancario [Tesis de "
    "pregrado, Universidad Nacional Mayor de San Marcos]. "
    "https://hdl.handle.net/20.500.12672/17840"
)
para(
    "Gómez Trapero, J. (2016). Desarrollo front-end y back-end de una aplicación "
    "Android distribuida para el sector de la restauración [Tesis de pregrado, "
    "Universidad Politécnica de Madrid]. https://oa.upm.es/43493/"
)
para(
    "Bedoya Alzate, E. (2021). Implementación de pruebas unitarias [Tesis de "
    "pregrado, Universidad de Antioquía]. https://hdl.handle.net/10495/20225"
)

doc.save(OUT)
print(f"Documento generado: {OUT}")
